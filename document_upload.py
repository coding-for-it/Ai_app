import streamlit as st
import pdfplumber
import docx
import easyocr
from PIL import Image
import fitz  # PyMuPDF
import os
import base64

# Cache EasyOCR reader to improve performance
@st.cache_resource
def load_easyocr_reader():
    return easyocr.Reader(['en'])

# Extract text from image using EasyOCR
def extract_text_image(file_path):
    reader = load_easyocr_reader()
    result = reader.readtext(file_path, detail=0)
    return ' '.join(result)

# Extract text from PDF using pdfplumber
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        pdf_text = ""
        for page in pdf.pages:
            pdf_text += page.extract_text() or ""
    return pdf_text

# Extract text from DOCX
def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

# Encode PDF to base64 for embedding
def encode_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode("utf-8")
    return f"data:application/pdf;base64,{base64_pdf}"

# Main Streamlit app
def document_upload_page():
    st.markdown("""
        <h1 style='text-align: center; color: #4F8BF9;'>📄 Document Upload and Viewer</h1>
    """, unsafe_allow_html=True)

    st.write("Upload a *PDF, **DOCX, or **Image* file to view its content.")

    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "png", "jpg", "jpeg"], help="Supported formats: .pdf, .docx, .png, .jpg, .jpeg")

    if uploaded_file is not None:
        st.success(f"File Uploaded Successfully: {uploaded_file.name}")

        file_path = os.path.join("temp", uploaded_file.name)
        os.makedirs("temp", exist_ok=True)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.read())

        # Display preview based on file type
        if uploaded_file.name.lower().endswith((".png", ".jpg", ".jpeg")):
            st.image(file_path, caption="Uploaded Image", use_container_width=True)
        elif uploaded_file.name.lower().endswith(".pdf"):
            st.subheader("PDF Preview")
            pdf_data = encode_pdf(file_path)
            st.markdown(f'<iframe src="{pdf_data}" width="100%" height="600px" style="border:none;"></iframe>', unsafe_allow_html=True)
        elif uploaded_file.name.lower().endswith(".docx"):
            st.subheader("DOCX Preview")
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                st.write(para.text)

        # Extract text based on file type
        if uploaded_file.name.endswith(".pdf"):
            extracted_text = extract_text_from_pdf(file_path)
        elif uploaded_file.name.endswith(".docx"):
            extracted_text = extract_text_from_docx(uploaded_file)
        elif uploaded_file.name.lower().endswith((".png", ".jpg", ".jpeg")):
            extracted_text = extract_text_image(file_path)
        else:
            st.error("Unsupported file format!")
            return

        # Display extracted text
        st.text_area("Extracted Text:", extracted_text, height=300, help="Scroll to read the full content.")

        # Download button
        st.download_button(
            label="💾 Download Extracted Text",
            data=extracted_text,
            file_name="extracted_text.txt",
            mime="text/plain"
        )

if __name__ == "__main__":
    document_upload_page()